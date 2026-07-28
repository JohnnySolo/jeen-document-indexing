"""Text extraction from PDF and DOCX files.

Responsibilities:
  - reject files that do not exist,
  - reject file types this module cannot read,
  - reject files that parse but yield no usable text,
  - return normalised plain text with paragraph breaks preserved.

Paragraph breaks are preserved deliberately: the paragraph chunking strategy
depends on them, so collapsing all whitespace would silently disable it.
"""

import re
from pathlib import Path

from errors import (
    ExtractionError,
    NoExtractableTextError,
    UnsupportedFileTypeError,
)
from config import SUPPORTED_EXTENSIONS

# Three or more consecutive newlines collapse to exactly two, which is the
# paragraph separator used downstream.
_EXCESS_BLANK_LINES = re.compile(r"\n{3,}")

# Runs of spaces and tabs collapse to a single space. Newlines are excluded from
# this class so paragraph structure survives.
_HORIZONTAL_WHITESPACE = re.compile(r"[ \t\u00a0]+")

# A word split across a line break by a hyphen, e.g. "regu-\nlation".
_HYPHEN_LINEBREAK = re.compile(r"(\w)-\n(\w)")


def clean_text(raw: str) -> str:
    """Normalise extracted text without destroying paragraph structure."""
    text = raw.replace("\r\n", "\n").replace("\r", "\n")
    text = _HYPHEN_LINEBREAK.sub(r"\1\2", text)
    text = _HORIZONTAL_WHITESPACE.sub(" ", text)
    text = "\n".join(line.strip() for line in text.split("\n"))
    text = _EXCESS_BLANK_LINES.sub("\n\n", text)
    return text.strip()


def _extract_pdf(path: Path) -> str:
    """Extract text from every page of a PDF."""
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ExtractionError(f"pypdf is not installed: {exc}") from exc

    try:
        reader = PdfReader(str(path))
        pages = [page.extract_text() or "" for page in reader.pages]
    except Exception as exc:
        raise ExtractionError(f"Could not read the PDF: {exc}") from exc

    return "\n\n".join(pages)


def _extract_docx(path: Path) -> str:
    """Extract text from DOCX paragraphs and table cells."""
    try:
        import docx
    except ImportError as exc:
        raise ExtractionError(f"python-docx is not installed: {exc}") from exc

    try:
        document = docx.Document(str(path))
    except Exception as exc:
        raise ExtractionError(f"Could not read the DOCX: {exc}") from exc

    parts = [p.text for p in document.paragraphs]

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n\n".join(parts)


def extract_text(file_path: str) -> str:
    """Return clean text from a PDF or DOCX file.

    Raises FileNotFoundError, UnsupportedFileTypeError, ExtractionError or
    NoExtractableTextError, each of which the CLI reports as a plain message.
    """
    path = Path(file_path).expanduser().resolve()

    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")
    if not path.is_file():
        raise FileNotFoundError(f"Not a file: {path}")

    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(SUPPORTED_EXTENSIONS)
        raise UnsupportedFileTypeError(
            f"Unsupported file type '{suffix or '(none)'}'. Supported: {supported}"
        )

    raw = _extract_pdf(path) if suffix == ".pdf" else _extract_docx(path)
    text = clean_text(raw)

    if not text:
        raise NoExtractableTextError(
            f"No extractable text found in {path.name}. "
            "If this is a scanned document it needs OCR first."
        )

    return text
